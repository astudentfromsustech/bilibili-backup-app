import pathlib as p
import re
import typing as t

import typing_extensions as te

from docx import api
from docx.shared import Pt, RGBColor
from docx.document import Document
from docx.text.font import Font
from pygments.lexers import get_lexer_by_name, guess_lexer, guess_lexer_for_filename
from pygments.styles import get_style_by_name


Path = t.Union[str, p.Path]
StrOrNone = t.Optional[str]


class Word:
    def __init__(self, root: Path = '.') -> None:
        self._root = p.Path(root).absolute()
        self._docx: Document = api.Document()
        self._compact = False
        self._page_break = False
        self._style = {}
        self._save = 'default.docx'

    def __enter__(self) -> te.Self:
        return self

    def __exit__(self, *args, **kwargs) -> None:
        self.save()

    @property
    def docx(self) -> Document:
        return self._docx

    def set_default(self) -> te.Self:
        self.set_docx_compact(False)
        self.set_docx_font(name='Times New Roman', size=Pt(7))
        self.set_docx_page_break(False)
        self.set_pygments_style('vs')
        return self

    def set_docx_compact(self, value: bool) -> te.Self:
        # self::compact
        self._compact = value
        return self

    def set_docx_font(self, **kwargs: t.Any) -> te.Self:
        # docx::font
        font: Font = self._docx.styles['Normal'].font
        for key, value in kwargs.items():
            setattr(font, key, value)
        return self

    def set_docx_page_break(self, value: bool) -> te.Self:
        # docx::page_break
        self._page_break = value
        return self

    def set_pygments_style(self, value: str) -> te.Self:
        # pygments::style
        self._style = dict(get_style_by_name(value))
        return self

    def save(self, path: t.Optional[Path] = None) -> te.Self:
        self._docx.save(path or self._save)
        return self

    def save_later(self, path: Path) -> te.Self:
        self._save = path
        return self

    def add_plain_by_text(self, text: str, title: StrOrNone = None) -> te.Self:
        if title is not None:
            self._docx.add_heading(title, 1)
        self._docx.add_paragraph(self._strip(text))
        if self._page_break:
            self._docx.add_page_break()
        return self

    def add_plain_by_path(self, path: Path, title: StrOrNone = None) -> te.Self:
        path_abs = self._path_abs(p.Path(path))
        if title is None:
            title = self._path_rel(path_abs).as_posix()
        return self.add_plain_by_text(path_abs.read_text(), title)

    def add_fancy_by_text(self, text: str, title: StrOrNone = None, lang: StrOrNone = None) -> te.Self:
        lexer = guess_lexer(text) if lang is None else get_lexer_by_name(lang)
        if title is not None:
            self._docx.add_heading(title, 1)
        paragraph = self._docx.add_paragraph()
        for type, value in lexer.get_tokens(self._strip(text)):
            style = self._style.get(type, {})
            color = style.get('color', None)
            run = paragraph.add_run(value)
            run.bold = style.get('bold', False)
            run.italic = style.get('italic', False)
            run.underline = style.get('underline', False)
            if color is not None:
                run.font.color.rgb = RGBColor.from_string(color)
        if self._page_break:
            self._docx.add_page_break()
        return self

    def add_fancy_by_path(self, path: Path, title: StrOrNone = None, lang: StrOrNone = None) -> te.Self:
        path_abs = self._path_abs(p.Path(path))
        text = path_abs.read_text()
        if title is None:
            title = self._path_rel(path_abs).as_posix()
        if lang is None:
            lang = guess_lexer_for_filename(path_abs.name, text).name
        return self.add_fancy_by_text(text, title, lang)

    def _path_abs(self, path: p.Path) -> p.Path:
        ans = p.Path(path)
        if ans.is_absolute():
            return ans
        else:
            return self._root / ans

    def _path_rel(self, path: p.Path) -> p.Path:
        try:
            return path.relative_to(self._root)
        except ValueError:
            return path

    def _strip(self, value: str) -> str:
        if self._compact:
            return re.sub(r'\n+', '\n', value.strip())
        else:
            return re.sub(r'\n{2,}', '\n\n', value.strip())


if __name__ == '__main__':
    import subprocess as sp

    root = p.Path(__file__).parents[1]

    with Word('.') \
            .set_default() \
            .set_docx_compact(False) \
            .set_pygments_style('friendly') \
            .save_later(root/'data'/'copyright.docx') as word:
        # files or directories
        files = ['app.py', 'script/copyright.py', 'script/word_cloud.py', 'Pipfile', 'Makefile'] + ['extra/app.py', 'extra/sync.py']
        directories = ['lib'] + ['extra/ks']
        for file in files:
            word.add_fancy_by_path(root/file)
        for directory in directories:
            for path in (root/directory).rglob('*'):
                if path.is_file():
                    word.add_fancy_by_path(path)
        # derivative
        sql = root / 'data' / 'db.sqlite'
        if sql.exists():
            cp = sp.run(['sqlite3', sql, '.dump video'], capture_output=True)
            word.add_fancy_by_text(cp.stdout.decode(), lang='sqlite3')
